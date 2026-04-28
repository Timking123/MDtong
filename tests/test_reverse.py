"""reverse 模块单元测试 — DOCX → Markdown 往返一致性"""

import os
import tempfile
import pytest
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from converter.reverse import docx_to_markdown
from converter.generator import convert


class TestDocxToMarkdown:
    def _make_docx(self, builder_fn):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            path = f.name
        doc = Document()
        builder_fn(doc)
        doc.save(path)
        return path

    def test_basic_paragraph(self):
        def build(doc):
            doc.add_paragraph('Hello World')
        path = self._make_docx(build)
        try:
            md = docx_to_markdown(path)
            assert 'Hello World' in md
        finally:
            os.unlink(path)

    def test_bold_text(self):
        def build(doc):
            p = doc.add_paragraph()
            run = p.add_run('加粗文本')
            run.bold = True
        path = self._make_docx(build)
        try:
            md = docx_to_markdown(path)
            assert '**加粗文本**' in md
        finally:
            os.unlink(path)

    def test_centered_title(self):
        def build(doc):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('文档标题')
            run.font.size = Pt(18)
        path = self._make_docx(build)
        try:
            md = docx_to_markdown(path)
            assert '# 文档标题' in md
        finally:
            os.unlink(path)

    def test_table(self):
        def build(doc):
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = '姓名'
            table.cell(0, 1).text = '部门'
            table.cell(1, 0).text = '张三'
            table.cell(1, 1).text = '技术部'
        path = self._make_docx(build)
        try:
            md = docx_to_markdown(path)
            assert '| 姓名 | 部门 |' in md
            assert '| 张三 | 技术部 |' in md
            assert '| --- | --- |' in md
        finally:
            os.unlink(path)

    def test_code_font_detection(self):
        def build(doc):
            p = doc.add_paragraph()
            run = p.add_run('var x = 1;')
            run.font.name = 'Consolas'
        path = self._make_docx(build)
        try:
            md = docx_to_markdown(path)
            assert '```' in md
        finally:
            os.unlink(path)

    def test_empty_document(self):
        def build(doc):
            pass
        path = self._make_docx(build)
        try:
            md = docx_to_markdown(path)
            assert md.strip() == '' or md.strip() == '\n'.strip()
        finally:
            os.unlink(path)

    def test_multiple_paragraphs(self):
        def build(doc):
            doc.add_paragraph('段落一')
            doc.add_paragraph('段落二')
            doc.add_paragraph('段落三')
        path = self._make_docx(build)
        try:
            md = docx_to_markdown(path)
            assert '段落一' in md
            assert '段落二' in md
            assert '段落三' in md
        finally:
            os.unlink(path)


class TestRoundTrip:
    def test_simple_roundtrip(self):
        md_input = '# 测试文档\n\n这是正文内容。'
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            docx_path = f.name
        try:
            convert(md_input, output_path=docx_path)
            md_output = docx_to_markdown(docx_path)
            assert '测试文档' in md_output
            assert '正文内容' in md_output
        finally:
            os.unlink(docx_path)

    def test_table_roundtrip(self):
        md_input = '# 报表\n\n| 项目 | 进度 |\n| --- | --- |\n| A | 80% |\n| B | 100% |'
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            docx_path = f.name
        try:
            convert(md_input, output_path=docx_path)
            md_output = docx_to_markdown(docx_path)
            assert '项目' in md_output
            assert '进度' in md_output
            assert '80%' in md_output
        finally:
            os.unlink(docx_path)
