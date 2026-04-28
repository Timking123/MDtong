"""generator 模块单元测试"""

import os
import tempfile
import pytest
from docx import Document

from converter.generator import (
    setup_document, convert, auto_filename,
    FORMAT_PRESETS, _resolve_format, _add_watermark,
)
from converter.parser import parse_markdown


class TestResolveFormat:
    def test_default(self):
        fmt = _resolve_format()
        assert fmt['title_font'] == '等线'
        assert fmt['body_size'] == 11

    def test_gongwen_preset(self):
        fmt = _resolve_format({'preset': '公文'})
        assert fmt['title_font'] == '方正小标宋体'
        assert fmt['body_font'] == '仿宋'
        assert fmt['body_size'] == 16

    def test_custom_overrides(self):
        fmt = _resolve_format({
            'preset': '自定义',
            'custom_overrides': {'body_size': 14, 'body_font': '宋体'},
        })
        assert fmt['body_size'] == 14
        assert fmt['body_font'] == '宋体'

    def test_none_overrides_ignored(self):
        fmt = _resolve_format({
            'preset': '自定义',
            'custom_overrides': {'body_size': None},
        })
        assert fmt['body_size'] == FORMAT_PRESETS['默认']['body_size']


class TestSetupDocument:
    def test_default_document(self):
        doc = setup_document()
        assert len(doc.sections) == 1

    def test_with_template(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            tpl_path = f.name
        try:
            tpl = Document()
            tpl.add_paragraph('模板内容')
            tpl.save(tpl_path)

            doc = setup_document(template_path=tpl_path)
            assert len(doc.paragraphs) == 0
        finally:
            os.unlink(tpl_path)

    def test_invalid_template_path(self):
        doc = setup_document(template_path='/nonexistent/template.docx')
        assert len(doc.sections) == 1


class TestConvert:
    def test_basic_conversion(self):
        md = '# 测试文档\n\n这是正文内容。'
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            out = f.name
        try:
            result = convert(md, output_path=out)
            assert os.path.isfile(result)
            doc = Document(result)
            assert len(doc.paragraphs) > 0
        finally:
            os.unlink(out)

    def test_with_table(self):
        md = '# 报表\n\n| A | B |\n| --- | --- |\n| 1 | 2 |'
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            out = f.name
        try:
            convert(md, output_path=out)
            doc = Document(out)
            assert len(doc.tables) == 1
        finally:
            os.unlink(out)

    def test_with_gongwen_preset(self):
        md = '# 公文标题\n\n一、第一段\n\n正文内容。'
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            out = f.name
        try:
            convert(md, output_path=out, format_settings={'preset': '公文'})
            doc = Document(out)
            assert len(doc.paragraphs) > 0
        finally:
            os.unlink(out)

    def test_empty_content_generates_default_title(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            out = f.name
        try:
            convert('', output_path=out)
            doc = Document(out)
            assert len(doc.paragraphs) > 0
        finally:
            os.unlink(out)

    def test_progress_callback(self):
        md = '# 标题\n\n内容'
        progress_values = []
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            out = f.name
        try:
            convert(md, output_path=out, progress_cb=lambda p, m: progress_values.append(p))
            assert 100 in progress_values
            assert len(progress_values) >= 3
        finally:
            os.unlink(out)

    def test_code_block(self):
        md = '# 文档\n\n```python\nprint("hello")\n```'
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            out = f.name
        try:
            convert(md, output_path=out)
            doc = Document(out)
            full_text = '\n'.join(p.text for p in doc.paragraphs)
            assert 'print("hello")' in full_text
        finally:
            os.unlink(out)


class TestAutoFilename:
    def test_generates_filename(self):
        blocks = [{'type': 'title', 'text': '测试文档'}]
        path = auto_filename(blocks, tempfile.gettempdir())
        assert '测试文档' in path
        assert path.endswith('.docx')


class TestWatermark:
    def test_xml_escape(self):
        doc = Document()
        _add_watermark(doc, '<script>alert("xss")</script>')
        xml = doc.sections[0].header._element.xml
        assert '<script>' not in xml
        assert '&lt;script&gt;' in xml

    def test_special_chars(self):
        doc = Document()
        _add_watermark(doc, '机密 & "内部"')
        xml = doc.sections[0].header._element.xml
        assert '&amp;' in xml
