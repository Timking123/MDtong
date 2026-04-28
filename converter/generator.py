"""DOCX 生成器 — 将结构化 block 列表转换为 Word 文档"""

import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE

from .parser import parse_inline_bold, parse_inline, parse_markdown

if getattr(sys, 'frozen', False):
    SCRIPT_DIR = sys._MEIPASS
else:
    SCRIPT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ── 格式预设 ──

FORMAT_PRESETS = {
    '默认': {
        'page_width_cm': 21.0,
        'page_height_cm': 29.7,
        'margin_top_cm': 2.54,
        'margin_bottom_cm': 2.54,
        'margin_left_cm': 3.18,
        'margin_right_cm': 3.18,
        'title_font': '等线',
        'title_font_western': 'Calibri',
        'title_size': 14,
        'heading_font': '等线',
        'heading_font_western': 'Calibri',
        'heading_size': 11,
        'body_font': '等线',
        'body_font_western': 'Calibri',
        'body_size': 11,
        'line_spacing_type': 'multiple',
        'line_spacing_value': 1.15,
        'first_line_indent_char': 0,
    },
    '公文': {
        'page_width_cm': 21.0,
        'page_height_cm': 29.7,
        'margin_top_cm': 3.7,
        'margin_bottom_cm': 3.5,
        'margin_left_cm': 2.8,
        'margin_right_cm': 2.6,
        'title_font': '方正小标宋体',
        'title_font_western': 'Times New Roman',
        'title_size': 22,
        'heading_font': '黑体',
        'heading_font_western': 'Times New Roman',
        'heading_size': 16,
        'body_font': '仿宋',
        'body_font_western': 'Times New Roman',
        'body_size': 16,
        'line_spacing_type': 'fixed',
        'line_spacing_value': 28.95,
        'first_line_indent_char': 2,
    },
}

DEFAULT_FORMAT = FORMAT_PRESETS['默认']


def _resolve_format(format_settings=None):
    """合并用户设置和预设，返回最终格式字典"""
    if not format_settings:
        return dict(DEFAULT_FORMAT)

    preset_name = format_settings.get('preset', '默认')
    base = dict(FORMAT_PRESETS.get(preset_name, DEFAULT_FORMAT))

    if preset_name == '自定义' or format_settings.get('custom_overrides'):
        overrides = format_settings.get('custom_overrides', format_settings)
        for key in base:
            if key in overrides and overrides[key] not in (None, ''):
                base[key] = overrides[key]

    return base


# ── 字体与段落设置工具函数 ──

def _set_run_fonts(run, east_asia, western, size=None, bold=False, color=None):
    run.font.name = western
    if size:
        run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), east_asia)


def _set_paragraph_spacing(p, fmt):
    pf = p.paragraph_format
    if fmt['line_spacing_type'] == 'fixed':
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(fmt['line_spacing_value'])
    elif fmt['line_spacing_type'] == 'multiple':
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = fmt['line_spacing_value']

    if fmt.get('first_line_indent_char', 0) > 0:
        indent_pt = fmt['body_size'] * fmt['first_line_indent_char']
        pf.first_line_indent = Pt(indent_pt)


def set_run_font(run, fmt=None):
    """兼容旧调用：使用 body 字体"""
    if fmt is None:
        fmt = DEFAULT_FORMAT
    _set_run_fonts(run, fmt['body_font'], fmt['body_font_western'])


def set_run_code_font(run):
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    rPr.append(shd)


def add_hyperlink(paragraph, text, url, fmt=None):
    if fmt is None:
        fmt = DEFAULT_FORMAT
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), fmt['body_font_western'])
    rFonts.set(qn('w:eastAsia'), fmt['body_font'])
    rPr.append(rFonts)
    run_el.append(rPr)
    run_text = OxmlElement('w:t')
    run_text.set(qn('xml:space'), 'preserve')
    run_text.text = text
    run_el.append(run_text)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def render_inline(paragraph, text, fmt=None):
    if fmt is None:
        fmt = DEFAULT_FORMAT
    parts = parse_inline(text)
    for part in parts:
        if part['type'] == 'text':
            run = paragraph.add_run(part['text'])
            if part.get('bold'):
                run.bold = True
            if part.get('strike'):
                run.font.strike = True
            set_run_font(run, fmt)
        elif part['type'] == 'code':
            run = paragraph.add_run(part['text'])
            set_run_code_font(run)
        elif part['type'] == 'link':
            add_hyperlink(paragraph, part['text'], part['url'], fmt)
        elif part['type'] == 'image':
            run = paragraph.add_run(f'[图片: {part["alt"]}]')
            set_run_font(run, fmt)


# ── 文档创建 ──

def setup_document(fmt=None, template_path=None):
    if fmt is None:
        fmt = DEFAULT_FORMAT
    if template_path and os.path.isfile(template_path):
        doc = Document(template_path)
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
        for t in list(doc.tables):
            t._element.getparent().remove(t._element)
    else:
        doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(fmt['page_width_cm'])
    section.page_height = Cm(fmt['page_height_cm'])
    section.left_margin = Cm(fmt['margin_left_cm'])
    section.right_margin = Cm(fmt['margin_right_cm'])
    section.top_margin = Cm(fmt['margin_top_cm'])
    section.bottom_margin = Cm(fmt['margin_bottom_cm'])

    style = doc.styles['Normal']
    style.font.name = fmt['body_font_western']
    style.font.size = Pt(fmt['body_size'])
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), fmt['body_font'])

    if fmt['line_spacing_type'] == 'fixed':
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(fmt['line_spacing_value'])
    elif fmt['line_spacing_type'] == 'multiple':
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = fmt['line_spacing_value']

    return doc


# ── Block 渲染函数 ──

def add_title(doc, text, fmt=None):
    if fmt is None:
        fmt = DEFAULT_FORMAT
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run_fonts(run, fmt['title_font'], fmt['title_font_western'],
                   size=fmt['title_size'], bold=True)
    _set_paragraph_spacing(p, fmt)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(4)


def add_section_header(doc, text, fmt=None):
    if fmt is None:
        fmt = DEFAULT_FORMAT
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_fonts(run, fmt['heading_font'], fmt['heading_font_western'],
                   size=fmt['heading_size'], bold=True)
    _set_paragraph_spacing(p, fmt)
    p.paragraph_format.first_line_indent = None


def add_sub_header(doc, text, fmt=None):
    if fmt is None:
        fmt = DEFAULT_FORMAT
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_fonts(run, fmt['heading_font'], fmt['heading_font_western'],
                   size=fmt['body_size'], bold=True)
    _set_paragraph_spacing(p, fmt)
    p.paragraph_format.first_line_indent = None


def add_label_content(doc, label, content, fmt=None):
    if fmt is None:
        fmt = DEFAULT_FORMAT
    p = doc.add_paragraph()
    run_label = p.add_run(f'{label}：')
    _set_run_fonts(run_label, fmt['body_font'], fmt['body_font_western'],
                   size=fmt['body_size'], bold=True)
    if content:
        render_inline(p, content, fmt)
    _set_paragraph_spacing(p, fmt)
    p.paragraph_format.first_line_indent = None


def add_body(doc, text, fmt=None):
    if fmt is None:
        fmt = DEFAULT_FORMAT
    p = doc.add_paragraph()
    render_inline(p, text, fmt)
    _set_paragraph_spacing(p, fmt)


def add_empty(doc):
    doc.add_paragraph()


def add_code_block(doc, code, language=''):
    for line in code.split('\n'):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), 'Consolas')


def add_image(doc, alt_text, img_path, base_dir=None):
    if base_dir and not os.path.isabs(img_path):
        img_path = os.path.join(base_dir, img_path)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.isfile(img_path):
        run = p.add_run()
        try:
            run.add_picture(img_path, width=Inches(5))
        except Exception:
            p.clear()
            run = p.add_run(f'[图片: {alt_text or img_path}]')
            set_run_font(run)
    else:
        run = p.add_run(f'[图片: {alt_text or img_path}]')
        set_run_font(run)

    if alt_text:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(alt_text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        set_run_font(run)


def add_task_item(doc, text, checked, fmt=None):
    if fmt is None:
        fmt = DEFAULT_FORMAT
    p = doc.add_paragraph()
    marker = '☑ ' if checked else '☐ '
    run = p.add_run(marker)
    set_run_font(run, fmt)
    render_inline(p, text, fmt)
    _set_paragraph_spacing(p, fmt)
    p.paragraph_format.first_line_indent = None


def add_list_item(doc, text, level, ordered, number=None, fmt=None):
    if fmt is None:
        fmt = DEFAULT_FORMAT
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25 * level)
    if ordered and number:
        prefix = f'{number}. '
    else:
        prefix = '• '
    run = p.add_run(prefix)
    set_run_font(run, fmt)
    render_inline(p, text, fmt)
    _set_paragraph_spacing(p, fmt)
    p.paragraph_format.first_line_indent = None


# ── 表格 ──

def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_cell_valign_center(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement('w:vAlign')
    v_align.set(qn('w:val'), 'center')
    tc_pr.append(v_align)


def set_table_properties(table):
    tbl_pr = table._tbl.tblPr
    cell_spacing = OxmlElement('w:tblCellSpacing')
    cell_spacing.set(qn('w:w'), '15')
    cell_spacing.set(qn('w:type'), 'dxa')
    tbl_pr.append(cell_spacing)
    cell_mar = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), '15')
        el.set(qn('w:type'), 'dxa')
        cell_mar.append(el)
    tbl_pr.append(cell_mar)


def add_table(doc, headers, rows, fmt=None):
    if fmt is None:
        fmt = DEFAULT_FORMAT
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = doc.styles['Normal Table']
    set_table_properties(table)

    for ci, header_text in enumerate(headers):
        cell = table.rows[0].cells[ci]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(header_text)
        _set_run_fonts(run, fmt['body_font'], fmt['body_font_western'],
                       size=fmt['body_size'], bold=True)
        set_cell_borders(cell)
        set_cell_valign_center(cell)

    for ri, row_data in enumerate(rows):
        padded = row_data + [''] * max(0, num_cols - len(row_data))
        for ci in range(num_cols):
            cell = table.rows[ri + 1].cells[ci]
            p = cell.paragraphs[0]
            p.clear()
            text = padded[ci] if ci < len(padded) else ''
            render_inline(p, text, fmt)
            set_cell_borders(cell)
            set_cell_valign_center(cell)


# ── 公文专用：红头分隔线 ──

def _add_red_line(doc):
    """添加公文红色分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '24')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'FF0000')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── 公文专用：页码格式 ── X ── ──

def _add_gongwen_page_number(doc):
    """公文页码格式：— X —（四号半角，居中）"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_pre = fp.add_run('— ')
    run_pre.font.size = Pt(14)
    run_pre.font.name = 'Times New Roman'

    run1 = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._element.append(fldChar1)

    run2 = fp.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)

    run3 = fp.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)

    run_suf = fp.add_run(' —')
    run_suf.font.size = Pt(14)
    run_suf.font.name = 'Times New Roman'


# ── 转换流水线 ──

def auto_filename(blocks, output_dir):
    title_block = next((b for b in blocks if b['type'] == 'title'), None)
    if title_block:
        name = title_block['text']
        name = re.sub(r'\s*[（(].+?[）)]', '', name)
        name = re.sub(r'[<>:"/\\|?*]', '', name)
        name = name.strip()
        if '会议纪要' not in name:
            name = f'{name}-会议纪要'
        filename = f'{name}.docx'
    else:
        filename = f'会议纪要-{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    return os.path.join(output_dir, filename)


def convert(text, output_path=None, progress_cb=None, doc_features=None,
            base_dir=None, format_settings=None, output_dir=None,
            cancel_event=None):
    def ensure_not_cancelled():
        if cancel_event is not None and cancel_event.is_set():
            raise InterruptedError('用户取消')

    ensure_not_cancelled()
    if progress_cb:
        progress_cb(5, '正在解析 Markdown...')

    fmt = _resolve_format(format_settings)
    blocks = parse_markdown(text)

    if not blocks or all(b['type'] == 'empty' for b in blocks):
        raise ValueError('未检测到有效的会议纪要内容')

    if progress_cb:
        progress_cb(15, f'解析完成，共 {len(blocks)} 个元素')
        progress_cb(20, '正在创建文档...')

    ensure_not_cancelled()
    doc = setup_document(fmt, template_path=(format_settings or {}).get('docx_template'))
    is_gongwen = (format_settings or {}).get('preset') == '公文'

    if doc_features:
        _apply_doc_features(doc, doc_features)
    elif is_gongwen:
        _add_gongwen_page_number(doc)

    total = len(blocks)
    title_done = False

    for i, block in enumerate(blocks):
        ensure_not_cancelled()
        t = block['type']

        # 公文模式：标题后自动插入红头分隔线
        if is_gongwen and t == 'title' and not title_done:
            add_title(doc, block['text'], fmt)
            _add_red_line(doc)
            title_done = True
        elif t == 'title':
            add_title(doc, block['text'], fmt)
        elif t == 'section_header':
            add_section_header(doc, block['text'], fmt)
        elif t == 'sub_header':
            add_sub_header(doc, block['text'], fmt)
        elif t == 'label_content':
            add_label_content(doc, block['label'], block['content'], fmt)
        elif t == 'body':
            add_body(doc, block['text'], fmt)
        elif t == 'empty':
            add_empty(doc)
        elif t == 'table':
            add_table(doc, block['headers'], block['rows'], fmt)
        elif t == 'code_block':
            add_code_block(doc, block['code'], block.get('language', ''))
        elif t == 'image':
            add_image(doc, block.get('alt', ''), block['path'], base_dir)
        elif t == 'task_item':
            add_task_item(doc, block['text'], block['checked'], fmt)
        elif t == 'list_item':
            add_list_item(
                doc, block['text'], block['level'],
                block.get('ordered', False), block.get('number'), fmt,
            )

        if progress_cb and total > 0:
            pct = 20 + int(70 * (i + 1) / total)
            progress_cb(pct, f'正在写入 {i + 1}/{total}...')

    if output_path is None:
        output_path = auto_filename(blocks, output_dir or SCRIPT_DIR)

    ensure_not_cancelled()
    if progress_cb:
        progress_cb(95, '正在保存文件...')

    doc.save(output_path)

    if progress_cb:
        progress_cb(100, '转换完成！')

    return output_path


def _apply_doc_features(doc, features):
    """应用文档专业功能（TOC、页眉页脚、水印等）"""
    section = doc.sections[0]

    if features.get('header_enabled') or features.get('logo_path'):
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if features.get('logo_path') and os.path.isfile(features['logo_path']):
            run = hp.add_run()
            try:
                run.add_picture(features['logo_path'], width=Inches(1.0))
            except Exception:
                pass

        header_text = features.get('header_text', '')
        if header_text:
            run = hp.add_run(f'  {header_text}')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            set_run_font(run)

    if features.get('page_number') or features.get('footer_enabled'):
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if features.get('page_number'):
            run = fp.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._element.append(fldChar1)

            run2 = fp.add_run()
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = ' PAGE '
            run2._element.append(instrText)

            run3 = fp.add_run()
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run3._element.append(fldChar2)

    if features.get('watermark_text'):
        _add_watermark(doc, features['watermark_text'])

    if features.get('toc_enabled'):
        _add_toc(doc)


def _add_toc(doc):
    """在文档开头插入目录占位域"""
    p = doc.add_paragraph()
    run = p.add_run('目录')
    run.bold = True
    run.font.size = Pt(14)
    set_run_font(run)

    p2 = doc.add_paragraph()
    run1 = p2.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._element.append(fldChar1)

    run2 = p2.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r' TOC \o "1-3" \h \z \u '
    run2._element.append(instrText)

    run3 = p2.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)

    doc.add_paragraph()


def _add_watermark(doc, text):
    """在页眉中添加文字水印"""
    from xml.sax.saxutils import escape
    safe_text = escape(text, {'"': '&quot;', "'": '&apos;'})

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pict = OxmlElement('w:pict')
    shape_xml = (
        f'<v:shape xmlns:v="urn:schemas-microsoft-com:vml" '
        f'xmlns:o="urn:schemas-microsoft-com:office:office" '
        f'style="position:absolute;margin-left:0;margin-top:0;'
        f'width:500pt;height:200pt;rotation:-45;z-index:-251657216" '
        f'fillcolor="silver" stroked="f" '
        f'type="#_x0000_t136">'
        f'<v:fill opacity=".25"/>'
        f'<v:textpath style="font-family:&quot;Microsoft YaHei&quot;;'
        f'font-size:60pt" string="{safe_text}"/>'
        f'<o:lock aspectratio="t"/>'
        f'</v:shape>'
    )

    from lxml import etree
    shape_el = etree.fromstring(shape_xml)
    pict.append(shape_el)

    run = OxmlElement('w:r')
    run.append(pict)
    p._p.append(run)
